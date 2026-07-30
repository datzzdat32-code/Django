"""
Script tự động tạo file Word (.docx) để nộp bài theo đúng yêu cầu:
Yêu cầu 4: Nộp link github cá nhân vào đây (tạo file word -> dán link github của dự án -> nộp file word vào phần "Nộp bài")
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_submission_doc(student_name="[Họ và Tên Sinh Viên]", student_id="[Mã Số Sinh Viên]", class_name="[Lớp Học Phần]", github_url="https://github.com/username/Chat-AI-Django"):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("BÁO CÁO NỘP BÀI TẬP BÀI THỰC HÀNH\nỨNG DỤNG WEB CHAT AI VỚI DJANGO & GOOGLE GEMINI")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 86, 219)
    
    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Môn học: Phát triển Ứng dụng Web / Trí tuệ Nhân tạo Tích hợp\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Information Box / Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    info_data = [
        ("Họ và tên sinh viên:", student_name),
        ("Mã số sinh viên (MSSV):", student_id),
        ("Lớp học phần:", class_name),
        ("Ngày nộp bài:", "30/07/2026")
    ]
    
    for i, (label, val) in enumerate(info_data):
        row_cells = table.rows[i].cells
        
        # Label cell
        p_label = row_cells[0].paragraphs[0]
        r_label = p_label.add_run(label)
        r_label.font.name = "Arial"
        r_label.font.bold = True
        r_label.font.size = Pt(11)
        
        # Value cell
        p_val = row_cells[1].paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.name = "Arial"
        r_val.font.size = Pt(11)
        r_val.font.color.rgb = RGBColor(30, 41, 59)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Heading 1: Link GitHub Repository
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. ĐƯỜNG DẪN DỰ ÁN GITHUB (GITHUB REPOSITORY LINK)")
    r_h1.font.name = "Arial"
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(26, 86, 219)
    
    p_link = doc.add_paragraph()
    p_link.paragraph_format.left_indent = Inches(0.2)
    r_link_label = p_link.add_run("👉 Link GitHub Repo: ")
    r_link_label.font.name = "Arial"
    r_link_label.font.bold = True
    r_link_label.font.size = Pt(12)
    
    r_link_val = p_link.add_run(github_url)
    r_link_val.font.name = "Arial"
    r_link_val.font.size = Pt(12)
    r_link_val.font.bold = True
    r_link_val.font.color.rgb = RGBColor(37, 99, 235)
    r_link_val.font.underline = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    
    # Heading 2: Tech Stack Summary
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. TÓM TẮT CÔNG NGHỆ VÀ TÍNH NĂNG ĐÃ THỰC HIỆN")
    r_h2.font.name = "Arial"
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(26, 86, 219)
    
    features = [
        "Backend: Django Framework (Python 3.14) tích hợp Google Gemini AI SDK (google-genai).",
        "Frontend: HTML5, Custom Vanilla CSS (Dark Mode Glassmorphism, Google Fonts, Responsive UI).",
        "Tích hợp AI: Hỗ trợ các model Gemini 2.5 Flash, Gemini 1.5 Flash, Gemini 2.5 Pro.",
        "Tính năng nổi bật: Chat realtime qua AJAX, giữ ngữ cảnh hội thoại theo Session, hỗ trợ Markdown rendering & Code highlighting với nút sao chép code nhanh.",
        "Cấu hình linh hoạt: Hỗ trợ cấu hình GEMINI_API_KEY trong file .env hoặc nhập trực tiếp API Key qua giao diện Web.",
        "Quản lý mã nguồn: Đã khởi tạo Git repository, cấu hình .gitignore loại bỏ file nhạy cảm (.env, db.sqlite3, .venv) và đẩy toàn bộ mã nguồn lên GitHub."
    ]
    
    for feat in features:
        p_feat = doc.add_paragraph(style='List Bullet')
        r_feat = p_feat.add_run(feat)
        r_feat.font.name = "Arial"
        r_feat.font.size = Pt(11)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Confirmation Footer
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_foot = p_foot.add_run("Sinh viên xác nhận đã hoàn thành bài tập và upload đầy đủ code lên GitHub.")
    r_foot.font.name = "Arial"
    r_foot.font.size = Pt(10)
    r_foot.font.italic = True
    
    output_filename = "Nop_Bai_Chat_AI.docx"
    doc.save(output_filename)
    print(f"[OK] Da tao thanh cong file Word nop bai: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_submission_doc()
